import base64
import io
import json
from datetime import datetime

from docx import Document
from docx.shared import Inches
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from sqlalchemy.orm import Session

from api.database import get_db
from api.models_db import AnalysisRun, EditHistory, Project, Upload

router = APIRouter(prefix="/report", tags=["report"])


def _get_run_or_404(run_id: int, db: Session):
    run = db.query(AnalysisRun).filter(AnalysisRun.id == run_id).first()
    if not run:
        raise HTTPException(status_code=404, detail="Analysis run not found")
    return run


def _build_context(run: AnalysisRun, db: Session):
    project = db.query(Project).filter(Project.id == run.project_id).first()
    upload = db.query(Upload).filter(Upload.project_id == run.project_id).first()
    result = json.loads(run.result_json or "{}")
    params = json.loads(run.parameters or "{}")
    flags = json.loads(upload.quality_flags if upload else "[]")
    return project, upload, result, params, flags


def _build_docx(run: AnalysisRun, db: Session) -> bytes:
    project, upload, result, params, flags = _build_context(run, db)

    doc = Document()
    doc.add_heading("QI Stat Studio Report", 0)
    if project:
        doc.add_heading(project.title or f"Project {project.id}", 1)

    doc.add_heading("Methods", 1)
    doc.add_paragraph(result.get("methods", "No methods description available."))

    doc.add_heading("Results", 1)
    doc.add_paragraph(result.get("result_summary", ""))

    tbl_data = result.get("table", [])
    if tbl_data:
        headers = list(tbl_data[0].keys())
        rt = doc.add_table(rows=1 + len(tbl_data), cols=len(headers))
        rt.style = "Table Grid"
        for j, h in enumerate(headers):
            rt.rows[0].cells[j].text = h.capitalize()
        for i, row in enumerate(tbl_data, 1):
            for j, h in enumerate(headers):
                rt.rows[i].cells[j].text = str(row.get(h, ""))

    # Figure
    fig_b64 = result.get("figure_base64")
    if fig_b64:
        img_data = base64.b64decode(fig_b64)
        doc.add_picture(io.BytesIO(img_data), width=Inches(5.5))

    # Interpretation placeholder (resident-edited text stored in edit_history)
    doc.add_heading("Interpretation", 1)
    doc.add_paragraph(result.get("interpretation", "[Resident interpretation will appear here after editing]"))

    # Limitations from DQ flags
    doc.add_heading("Limitations", 1)
    if flags:
        for f in flags:
            doc.add_paragraph(f"• {f.get('msg', str(f))}", style="List Bullet")
    else:
        doc.add_paragraph("No data quality issues were flagged for this dataset.")

    # Code supplement
    if run.code_r:
        doc.add_heading("Statistical Code Supplement (R)", 1)
        doc.add_paragraph(run.code_r, style="No Spacing")
    if run.code_spss:
        doc.add_heading("Statistical Code Supplement (SPSS)", 1)
        doc.add_paragraph(run.code_spss, style="No Spacing")
    if run.code_sas:
        doc.add_heading("Statistical Code Supplement (SAS)", 1)
        doc.add_paragraph(run.code_sas, style="No Spacing")

    # Audit trail
    doc.add_heading("Audit Trail", 1)
    audit_rows = [
        ["Template", run.template],
        ["Parameters", json.dumps(params, indent=2)],
        ["Source file", upload.filename if upload else "N/A"],
        ["Run date", datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")],
    ]
    tbl = doc.add_table(rows=len(audit_rows), cols=2)
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(audit_rows):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = str(v)

    edits = db.query(EditHistory).filter(EditHistory.project_id == run.project_id).all()
    if edits:
        doc.add_heading("Resident Edits", 2)
        edit_tbl = doc.add_table(rows=1 + len(edits), cols=3)
        edit_tbl.style = "Table Grid"
        for j, header in enumerate(["Field", "Original Text", "Edited Text"]):
            edit_tbl.rows[0].cells[j].text = header
        for i, e in enumerate(edits, 1):
            edit_tbl.rows[i].cells[0].text = e.field or ""
            edit_tbl.rows[i].cells[1].text = e.original_text or ""
            edit_tbl.rows[i].cells[2].text = e.edited_text or ""

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_pdf(run: AnalysisRun, db: Session) -> bytes:
    project, upload, result, params, flags = _build_context(run, db)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    title = project.title if project else f"Project {run.project_id}"
    story.append(Paragraph("QI Stat Studio Report", styles["Title"]))
    story.append(Paragraph(title, styles["Heading1"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Methods", styles["Heading2"]))
    story.append(Paragraph(result.get("methods", ""), styles["Normal"]))
    story.append(Spacer(1, 8))

    story.append(Paragraph("Results", styles["Heading2"]))
    story.append(Paragraph(result.get("result_summary", ""), styles["Normal"]))
    story.append(Spacer(1, 8))

    tbl_data = result.get("table", [])
    if tbl_data:
        headers = list(tbl_data[0].keys())
        rows = [[h.capitalize() for h in headers]] + [
            [str(r.get(h, "")) for h in headers] for r in tbl_data
        ]
        col_w = 470 / len(headers)
        pt = Table(rows, colWidths=[col_w] * len(headers))
        pt.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ]))
        story.append(pt)
        story.append(Spacer(1, 8))

    fig_b64 = result.get("figure_base64")
    if fig_b64:
        img_data = base64.b64decode(fig_b64)
        rl_img = RLImage(io.BytesIO(img_data), width=400, height=200)
        story.append(rl_img)
        story.append(Spacer(1, 8))

    story.append(Paragraph("Interpretation", styles["Heading2"]))
    story.append(Paragraph(result.get("interpretation", "[Resident interpretation pending]"), styles["Normal"]))
    story.append(Spacer(1, 8))

    story.append(Paragraph("Limitations", styles["Heading2"]))
    if flags:
        for f in flags:
            story.append(Paragraph(f"• {f.get('msg', str(f))}", styles["Normal"]))
    else:
        story.append(Paragraph("No data quality issues flagged.", styles["Normal"]))
    story.append(Spacer(1, 8))

    if run.code_r:
        story.append(Paragraph("Statistical Code Supplement (R)", styles["Heading2"]))
        story.append(Paragraph(run.code_r, styles["Code"]))
        story.append(Spacer(1, 8))
    if run.code_spss:
        story.append(Paragraph("Statistical Code Supplement (SPSS)", styles["Heading2"]))
        story.append(Paragraph(run.code_spss, styles["Code"]))
        story.append(Spacer(1, 8))
    if run.code_sas:
        story.append(Paragraph("Statistical Code Supplement (SAS)", styles["Heading2"]))
        story.append(Paragraph(run.code_sas, styles["Code"]))
        story.append(Spacer(1, 8))

    story.append(Paragraph("Audit Trail", styles["Heading2"]))
    audit_data = [
        ["Template", run.template],
        ["Source file", upload.filename if upload else "N/A"],
        ["Run date", datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")],
    ]
    t = Table(audit_data, colWidths=[120, 350])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
    ]))
    story.append(t)

    edits = db.query(EditHistory).filter(EditHistory.project_id == run.project_id).all()
    if edits:
        story.append(Spacer(1, 8))
        story.append(Paragraph("Resident Edits", styles["Heading2"]))
        edit_data = [["Field", "Original Text", "Edited Text"]] + [
            [e.field or "", e.original_text or "", e.edited_text or ""] for e in edits
        ]
        et = Table(edit_data, colWidths=[80, 220, 170])
        et.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey)]))
        story.append(et)

    doc.build(story)
    return buf.getvalue()


@router.get("/{run_id}/docx")
def download_docx(run_id: int, db: Session = Depends(get_db)):
    run = _get_run_or_404(run_id, db)
    content = _build_docx(run, db)
    return StreamingResponse(
        io.BytesIO(content),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=qi_report_{run_id}.docx"},
    )


@router.get("/{run_id}/pdf")
def download_pdf(run_id: int, db: Session = Depends(get_db)):
    run = _get_run_or_404(run_id, db)
    content = _build_pdf(run, db)
    return StreamingResponse(
        io.BytesIO(content),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=qi_report_{run_id}.pdf"},
    )
