"""Tests for /report/{id}/docx and /report/{id}/pdf endpoints — bytes and content."""
import io, json, os
os.environ.setdefault("FERNET_KEY", "AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA=")

from docx import Document
from fastapi.testclient import TestClient
from api.main import app
from api.database import engine, Base, SessionLocal
from api.models_db import Project, AnalysisRun, Upload

Base.metadata.create_all(bind=engine)
client = TestClient(app)


def _seed_run(template="run_chart", dq_flags=None):
    db = SessionLocal()
    p = Project(title="Test Project", description="test")
    db.add(p); db.flush()

    u = Upload(
        project_id=p.id, filename="data.csv",
        encrypted_path="/tmp/fake.enc",
        quality_flags=json.dumps(dq_flags or []),
    )
    db.add(u); db.flush()

    result = {"methods": "A run chart was used.", "result_summary": "Median=5.0", "figure_base64": None}
    run = AnalysisRun(
        project_id=p.id, template=template,
        parameters=json.dumps({"date_col": "encounter_date", "value_col": "hba1c"}),
        result_json=json.dumps(result), code_r="# R code here",
    )
    db.add(run); db.commit()
    run_id = run.id
    db.close()
    return run_id


def test_docx_returns_bytes():
    rid = _seed_run()
    resp = client.get(f"/report/{rid}/docx")
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith("application/vnd.openxmlformats")
    assert len(resp.content) > 1000  # non-empty Word doc


def test_pdf_returns_bytes():
    rid = _seed_run()
    resp = client.get(f"/report/{rid}/pdf")
    assert resp.status_code == 200
    assert resp.headers["content-type"] == "application/pdf"
    assert resp.content[:4] == b"%PDF"


def test_report_404():
    resp = client.get("/report/99999/docx")
    assert resp.status_code == 404


# ── content tests ──────────────────────────────────────────────────────────

def _docx_text(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def test_docx_contains_methods_section():
    rid = _seed_run()
    resp = client.get(f"/report/{rid}/docx")
    text = _docx_text(resp.content)
    assert "Methods" in text


def test_docx_contains_audit_trail():
    rid = _seed_run()
    resp = client.get(f"/report/{rid}/docx")
    text = _docx_text(resp.content)
    assert "Audit Trail" in text
    assert "run_chart" in text   # template name present


def test_docx_contains_limitations_section():
    rid = _seed_run()
    resp = client.get(f"/report/{rid}/docx")
    text = _docx_text(resp.content)
    assert "Limitations" in text


def test_docx_limitations_lists_dq_flags():
    rid = _seed_run(dq_flags=[
        {"col": "hba1c", "rule": "check_missing", "severity": "WARNING", "msg": "hba1c is 15% missing"}
    ])
    resp = client.get(f"/report/{rid}/docx")
    text = _docx_text(resp.content)
    assert "hba1c is 15% missing" in text


def test_docx_no_flags_says_no_issues():
    rid = _seed_run(dq_flags=[])
    resp = client.get(f"/report/{rid}/docx")
    text = _docx_text(resp.content)
    assert "No data quality issues" in text


def test_docx_contains_r_code_supplement():
    rid = _seed_run()
    resp = client.get(f"/report/{rid}/docx")
    text = _docx_text(resp.content)
    assert "R code here" in text


def test_pdf_contains_audit_trail_text():
    """PDF should contain the template name in its text content."""
    import pypdf
    rid = _seed_run()
    resp = client.get(f"/report/{rid}/pdf")
    reader = pypdf.PdfReader(io.BytesIO(resp.content))
    all_text = "".join(page.extract_text() or "" for page in reader.pages)
    assert "run_chart" in all_text
    assert "Audit Trail" in all_text
